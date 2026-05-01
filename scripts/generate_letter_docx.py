"""Render the markdown letter drafts under ``docs/letters/`` as
ready-to-sign Microsoft Word ``.docx`` files in ``docs/letters/word/``.

This is the project's reproducible converter from the source-tracked
markdown drafts to the Word format the project lead prints, signs,
and sends. The script is deterministic — re-running with the same
content produces byte-equivalent ``.docx`` output (modulo timestamps
inside the Word XML, which python-docx fills in automatically).

The Word output contains ONLY the letter content (sender block, date
placeholder, recipient block, bold subject line, salutation, body
with numbered lists, sign-off, signature block). The markdown drafts'
meta-content (status header, project context, action items, suggested
addressee section) does NOT appear in the Word output — those are
internal planning notes.

Two files are produced for the APH letter (one for the House, one
for the Senate) because the markdown draft is a single template
with bracketed alternatives. The AEC GIS letter produces a single
file.

Run from the project root with python-docx available:

    python3 -m venv /tmp/docx_venv
    /tmp/docx_venv/bin/pip install python-docx
    /tmp/docx_venv/bin/python scripts/generate_letter_docx.py

Or, if python-docx is later added to the backend lockfile (it is NOT
today), run with the backend venv directly.
"""

from __future__ import annotations

from pathlib import Path
from typing import Final

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


# Resolve relative to the script so re-runs from any cwd are stable.
PROJECT_ROOT: Final[Path] = Path(__file__).resolve().parent.parent
OUTPUT_DIR: Final[Path] = PROJECT_ROOT / "docs" / "letters" / "word"


SENDER_BLOCK: Final[list[str]] = [
    "[Project Lead Name]",
    "Australian Political Influence Transparency Project",
    "[Street Address]",
    "[Suburb] [State] [Postcode]",
    "[Email] · [Phone]",
]


def _new_document() -> Document:
    """Create a Document configured for a clean Australian business letter.

    Calibri 11pt body, 2.5 cm margins all round, single-spaced paragraphs
    with explicit before/after spacing. These are the defaults
    Australian Government correspondence usually uses.
    """
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    return doc


def _add_paragraph(
    doc: Document,
    text: str = "",
    *,
    bold: bool = False,
    align: int | None = None,
    space_before: float | None = None,
    space_after: float | None = None,
) -> None:
    paragraph = doc.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    if space_before is not None:
        paragraph.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        paragraph.paragraph_format.space_after = Pt(space_after)
    if text:
        run = paragraph.add_run(text)
        run.bold = bold


def _add_block(
    doc: Document,
    lines: list[str],
    *,
    align: int | None = None,
) -> None:
    """Add a multi-line block as one paragraph with hard line breaks."""
    paragraph = doc.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(6)
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        if index < len(lines) - 1:
            run.add_break()


BodyEntry = str | tuple[str, list[str]]


def _build_letter(
    *,
    sender_lines: list[str],
    recipient_lines: list[str],
    subject: str,
    salutation: str,
    body_paragraphs: list[BodyEntry],
    closing: str,
    signature_lines: list[str],
) -> Document:
    doc = _new_document()

    # Sender block, right-aligned at top.
    _add_block(doc, sender_lines, align=WD_ALIGN_PARAGRAPH.RIGHT)

    # Date placeholder, right-aligned.
    _add_paragraph(
        doc,
        "[Date]",
        align=WD_ALIGN_PARAGRAPH.RIGHT,
        space_before=6,
        space_after=18,
    )

    # Recipient block, left-aligned.
    _add_block(doc, recipient_lines)

    # Subject line, bold, with extra space.
    _add_paragraph(
        doc,
        subject,
        bold=True,
        space_before=12,
        space_after=12,
    )

    # Salutation.
    _add_paragraph(doc, salutation, space_after=6)

    # Body — strings are plain paragraphs, tuples produce numbered lists.
    for entry in body_paragraphs:
        if isinstance(entry, str):
            _add_paragraph(doc, entry, space_after=10)
        else:
            lead, items = entry
            if lead:
                _add_paragraph(doc, lead, space_after=4)
            for item_index, item in enumerate(items):
                paragraph = doc.add_paragraph(style="List Number")
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.add_run(item)
                if item_index == len(items) - 1:
                    paragraph.paragraph_format.space_after = Pt(10)

    # Closing.
    _add_paragraph(doc, closing, space_before=6, space_after=36)

    # Signature block.
    _add_block(doc, signature_lines)

    return doc


# ----------------------------------------------------------------------
# APH letter — House version
# ----------------------------------------------------------------------

APH_HOUSE_RECIPIENT: Final[list[str]] = [
    "The Clerk of the House of Representatives",
    "Department of the House of Representatives",
    "Parliament House",
    "Canberra ACT 2600",
]

APH_SUBJECT: Final[str] = (
    "Re: Public-redistribution exception request — APH website material "
    "(CC BY-NC-ND 4.0)"
)

APH_HOUSE_BODY: Final[list[BodyEntry]] = [
    "I am the lead on the Australian Political Influence Transparency "
    "project, a public-interest, source-backed transparency tool that "
    "publishes a reproducible link from disclosed federal political "
    "records back to the original public source documents.",

    "The project's data layer parses material that the Department of the "
    "House of Representatives publishes under Creative Commons "
    "Attribution-NonCommercial-NoDerivatives 4.0 International on the "
    "Parliament of Australia website. Specifically, we ingest:",

    (
        "",
        [
            "the Register of Members' Interests PDFs;",
            "the Members contacts CSV from "
            "https://www.aph.gov.au/Senators_and_Members/Members/...; and",
            "House Votes and Proceedings decision-record indexes and PDFs.",
        ],
    ),

    "Our public app surfaces these records alongside the original PDFs and "
    "clearly labels every claim with its evidence tier and attribution "
    "caveat. We do not paywall the data, do not run advertising on it, "
    "and are not a commercial entity. We treat structured-record "
    "transformations of the source PDFs as derivative works for the "
    "purposes of the CC BY-NC-ND 4.0 licence, and we are not yet "
    "publishing those transformations.",

    "Could the Department please advise:",

    (
        "",
        [
            "Whether the Department considers our structured-record "
            "transformation of the Register of Members' Interests, the "
            "Members contacts CSV, and House Votes and Proceedings to be "
            "a derivative work for the purposes of the CC BY-NC-ND 4.0 "
            "licence;",
            "If so, whether the Department is willing to grant an "
            "exception permitting us to publish those structured "
            "records — with attribution and any other reasonable "
            "conditions the Department asks for; and",
            "If an exception is not appropriate, what surface (e.g. "
            "verbatim PDFs and CSVs with no JSON or CSV transformation) "
            "would be permissible without an exception.",
        ],
    ),

    "Our reproducibility policy is at docs/reproducibility.md in the "
    "project repository. Every public claim travels with its evidence "
    "tier, attribution caveat, and a link back to the original source "
    "document. We would welcome any conditions the Department considers "
    "appropriate.",
]

APH_SIG: Final[list[str]] = [
    "Yours faithfully,",
    "",
    "",
    "",
    "[Project Lead Name]",
    "Project Lead, Australian Political Influence Transparency Project",
]


# ----------------------------------------------------------------------
# APH letter — Senate version
# ----------------------------------------------------------------------

APH_SENATE_RECIPIENT: Final[list[str]] = [
    "The Clerk of the Senate",
    "Department of the Senate",
    "Parliament House",
    "Canberra ACT 2600",
]

APH_SENATE_BODY: Final[list[BodyEntry]] = [
    "I am the lead on the Australian Political Influence Transparency "
    "project, a public-interest, source-backed transparency tool that "
    "publishes a reproducible link from disclosed federal political "
    "records back to the original public source documents.",

    "The project's data layer parses material that the Department of the "
    "Senate publishes under Creative Commons Attribution-NonCommercial-"
    "NoDerivatives 4.0 International on the Parliament of Australia "
    "website. Specifically, we ingest:",

    (
        "",
        [
            "the Register of Senators' Interests PDFs;",
            "the Senators contacts CSV from "
            "https://www.aph.gov.au/Senators_and_Members/Senators/...; and",
            "Senate Journals decision-record indexes and PDFs.",
        ],
    ),

    "Our public app surfaces these records alongside the original PDFs and "
    "clearly labels every claim with its evidence tier and attribution "
    "caveat. We do not paywall the data, do not run advertising on it, "
    "and are not a commercial entity. We treat structured-record "
    "transformations of the source PDFs as derivative works for the "
    "purposes of the CC BY-NC-ND 4.0 licence, and we are not yet "
    "publishing those transformations.",

    "Could the Department please advise:",

    (
        "",
        [
            "Whether the Department considers our structured-record "
            "transformation of the Register of Senators' Interests, the "
            "Senators contacts CSV, and Senate Journals to be a "
            "derivative work for the purposes of the CC BY-NC-ND 4.0 "
            "licence;",
            "If so, whether the Department is willing to grant an "
            "exception permitting us to publish those structured "
            "records — with attribution and any other reasonable "
            "conditions the Department asks for; and",
            "If an exception is not appropriate, what surface (e.g. "
            "verbatim PDFs and CSVs with no JSON or CSV transformation) "
            "would be permissible without an exception.",
        ],
    ),

    "Our reproducibility policy is at docs/reproducibility.md in the "
    "project repository. Every public claim travels with its evidence "
    "tier, attribution caveat, and a link back to the original source "
    "document. We would welcome any conditions the Department considers "
    "appropriate.",
]


# ----------------------------------------------------------------------
# AEC GIS letter
# ----------------------------------------------------------------------

AEC_GIS_RECIPIENT: Final[list[str]] = [
    "AEC GIS Section / AEC Electorate Mapping",
    "Australian Electoral Commission",
    "Locked Bag 4007",
    "Canberra ACT 2601",
]

AEC_GIS_SUBJECT: Final[str] = (
    "Re: Public-redistribution clarification — AEC GIS data "
    "(End-user Licence, Derivative Product permission)"
)

AEC_GIS_BODY: Final[list[BodyEntry]] = [
    "I am the lead on the Australian Political Influence Transparency "
    "project, a public-interest, source-backed transparency tool that "
    "publishes a reproducible link from disclosed federal political "
    "records back to the original public source documents.",

    "The project ingests the AEC's federal electorate boundary "
    "shapefile and serves it as a runtime map layer in the project's "
    "public app. Specifically:",

    (
        "",
        [
            "we download the official current national federal "
            "electorate ESRI shapefile from "
            "https://www.aec.gov.au/electorates/gis/...;",
            "we re-project the geometry from GDA94 to EPSG:4326 "
            "(GeoJSON / PostGIS) and serve it as vector tiles from a "
            "server the project operates;",
            "the public app renders those tiles with the verbatim "
            "attribution \"© Commonwealth of Australia "
            "(Australian Electoral Commission) [year]\", linked to the "
            "AEC GIS licence page; and",
            "the AEC's official boundary geometry is unchanged in "
            "storage (the only display-time variation is a coastline "
            "clip layered on top, sourced separately).",
        ],
    ),

    "We treat this combined treatment as a Derivative Product within the "
    "AEC's GIS End-user Licence permission. Could the AEC please confirm:",

    (
        "",
        [
            "Whether the AEC considers our re-projected, vector-tiled, "
            "publicly-served treatment of the boundary geometry a "
            "Derivative Product within the meaning of the End-user "
            "Licence;",
            "Whether the AEC's standard attribution string suffices, or "
            "whether the AEC would like an additional or different "
            "attribution form; and",
            "Whether the AEC has any other conditions the project "
            "should apply before public launch (we explicitly preserve "
            "the warranty disclaimer the licence states; we do not "
            "assert AEC endorsement in any UI surface; and we link the "
            "original AEC source documents alongside the rendered map).",
        ],
    ),

    "Our reproducibility policy is at docs/reproducibility.md in the "
    "project repository, and the per-source licence audit is at "
    "docs/source_licences.md. We would welcome any conditions the AEC "
    "considers appropriate.",
]


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    aph_house = _build_letter(
        sender_lines=SENDER_BLOCK,
        recipient_lines=APH_HOUSE_RECIPIENT,
        subject=APH_SUBJECT,
        salutation="Dear Clerk,",
        body_paragraphs=APH_HOUSE_BODY,
        closing="Yours faithfully,",
        signature_lines=APH_SIG,
    )
    aph_house_path = OUTPUT_DIR / "aph_public_redistribution_request_house.docx"
    aph_house.save(aph_house_path)
    print(f"wrote {aph_house_path.relative_to(PROJECT_ROOT)} "
          f"({aph_house_path.stat().st_size} bytes)")

    aph_senate = _build_letter(
        sender_lines=SENDER_BLOCK,
        recipient_lines=APH_SENATE_RECIPIENT,
        subject=APH_SUBJECT,
        salutation="Dear Clerk,",
        body_paragraphs=APH_SENATE_BODY,
        closing="Yours faithfully,",
        signature_lines=APH_SIG,
    )
    aph_senate_path = OUTPUT_DIR / "aph_public_redistribution_request_senate.docx"
    aph_senate.save(aph_senate_path)
    print(f"wrote {aph_senate_path.relative_to(PROJECT_ROOT)} "
          f"({aph_senate_path.stat().st_size} bytes)")

    aec_gis = _build_letter(
        sender_lines=SENDER_BLOCK,
        recipient_lines=AEC_GIS_RECIPIENT,
        subject=AEC_GIS_SUBJECT,
        salutation="Dear Sir / Madam,",
        body_paragraphs=AEC_GIS_BODY,
        closing="Yours faithfully,",
        signature_lines=APH_SIG,
    )
    aec_gis_path = OUTPUT_DIR / "aec_gis_public_redistribution_request.docx"
    aec_gis.save(aec_gis_path)
    print(f"wrote {aec_gis_path.relative_to(PROJECT_ROOT)} "
          f"({aec_gis_path.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
